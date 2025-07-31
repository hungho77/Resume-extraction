import os
import re
import json
from typing import Dict, List, Any, Optional
from pathlib import Path
import logging

import torch
import pandas as pd
from docling.document_converter import DocumentConverter
from docx import Document
from PIL import Image

# Fallback PDF processing
try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

# OCR processing
try:
    import easyocr
    import cv2
    import numpy as np
    from PIL import Image
    EASYOCR_AVAILABLE = True
except ImportError:
    EASYOCR_AVAILABLE = False

# SmolDocLing imports
try:
    import torch
    from transformers import AutoProcessor, AutoModelForVision2Seq
    from transformers.image_utils import load_image
    from docling_core.types.doc import DoclingDocument
    from docling_core.types.doc.document import DocTagsDocument
    SMOLDOCLING_AVAILABLE = True
except ImportError:
    SMOLDOCLING_AVAILABLE = False

from .config import docling_config

logger = logging.getLogger(__name__)

class ResumeDocumentProcessor:
    """Enhanced document processor for resume parsing using DocLing with OCR"""
    
    def __init__(self):
        self.converter = DocumentConverter()
        self.use_ocr = docling_config.use_ocr
        self.ocr_language = docling_config.ocr_language
        
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using DocLing with OCR capabilities"""
        text = ""
        
        try:
            # Use DocLing for PDF processing with OCR
            logger.info("Using DocLing for PDF extraction with OCR")
            
            # Convert document with DocLing
            conversion_result = self.converter.convert(file_path)
            
            # Extract text from DocLing result
            if hasattr(conversion_result, 'text'):
                text = conversion_result.text
            elif hasattr(conversion_result, 'content'):
                text = conversion_result.content
            elif hasattr(conversion_result, 'pages'):
                # Extract text from all pages
                text = "\n".join([page.text for page in conversion_result.pages if hasattr(page, 'text')])
            elif isinstance(conversion_result, dict):
                text = conversion_result.get('text', '')
            else:
                # Fallback: try to get text from the document object
                text = str(conversion_result)
            
            logger.info("DocLing PDF extraction successful")
            
        except Exception as e:
            logger.error(f"DocLing PDF extraction failed: {e}")
            text = ""
        
        # If DocLing failed or returned empty text, try fallback
        if not text.strip():
            logger.info("DocLing returned empty text, trying fallback PDF extraction with pypdf...")
            text = self.extract_text_from_pdf_fallback(file_path)
            if text:
                logger.info("Fallback PDF extraction successful")
            else:
                logger.info("pypdf fallback failed, trying EasyOCR extraction...")
                text = self.extract_text_from_pdf_ocr(file_path)
                if text:
                    logger.info("EasyOCR PDF extraction successful")
                else:
                    logger.info("EasyOCR failed, trying SmolDocLing extraction...")
                    text = self.extract_text_from_pdf_smoldocling(file_path)
                    if text:
                        logger.info("SmolDocLing PDF extraction successful")
                    else:
                        logger.error("All PDF extraction methods failed")
        
        return text.strip()
    
    def extract_text_from_pdf_fallback(self, file_path: str) -> str:
        """Fallback PDF text extraction using pypdf"""
        if not PYPDF_AVAILABLE:
            return ""
        
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Fallback PDF extraction failed: {e}")
            return ""
    
    def extract_text_from_pdf_ocr(self, file_path: str) -> str:
        """Extract text from PDF using OCR (EasyOCR)"""
        if not EASYOCR_AVAILABLE:
            return ""
        
        try:
            import fitz  # PyMuPDF
            import easyocr
            
            # Initialize EasyOCR
            reader = easyocr.Reader(['en'])
            
            # Open PDF with PyMuPDF
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Convert page to image
                pix = page.get_pixmap()
                img_data = pix.tobytes("png")
                
                # Convert to OpenCV format
                nparr = np.frombuffer(img_data, np.uint8)
                img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                
                # Perform OCR
                results = reader.readtext(img)
                
                # Extract text from results
                page_text = " ".join([result[1] for result in results])
                text += page_text + "\n"
            
            doc.close()
            return text.strip()
            
        except ImportError:
            logger.warning("PyMuPDF not available for OCR processing")
            return ""
        except Exception as e:
            logger.error(f"OCR PDF extraction failed: {e}")
            return ""
    
    def extract_text_from_pdf_smoldocling(self, file_path: str) -> str:
        """Extract text from PDF using SmolDocLing OCR"""
        if not SMOLDOCLING_AVAILABLE:
            return ""
        
        try:
            import fitz  # PyMuPDF
            
            # Initialize SmolDocLing
            model_name = "ds4sd/SmolDocling-256M-preview"
            device = "cuda" if torch.cuda.is_available() else "cpu"
            
            processor = AutoProcessor.from_pretrained(model_name)
            model = AutoModelForVision2Seq.from_pretrained(
                model_name,
                torch_dtype=torch.bfloat16,
                _attn_implementation="eager",  # Use eager to avoid FlashAttention2 dependency
            ).to(device)
            
            # Open PDF with PyMuPDF
            doc = fitz.open(file_path)
            all_text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Convert page to image
                pix = page.get_pixmap()
                img_data = pix.tobytes("png")
                
                # Save temporary image
                temp_image_path = f"/tmp/smoldocling_page_{page_num}.png"
                with open(temp_image_path, "wb") as f:
                    f.write(img_data)
                
                # Load image
                image = load_image(temp_image_path)
                
                # Create input messages
                messages = [
                    {
                        "role": "user",
                        "content": [
                            {"type": "image"},
                            {"type": "text", "text": "Convert this page to docling."}
                        ]
                    },
                ]
                
                # Prepare inputs
                prompt = processor.apply_chat_template(messages, add_generation_prompt=True)
                inputs = processor(text=prompt, images=[image], return_tensors="pt")
                inputs = inputs.to(device)
                
                # Generate outputs
                generated_ids = model.generate(**inputs, max_new_tokens=8192)
                prompt_length = inputs.input_ids.shape[1]
                trimmed_generated_ids = generated_ids[:, prompt_length:]
                doctags = processor.batch_decode(
                    trimmed_generated_ids,
                    skip_special_tokens=False,
                )[0].lstrip()
                
                # Convert to DoclingDocument and extract text
                doctags_doc = DocTagsDocument.from_doctags_and_image_pairs([doctags], [image])
                doc_ling = DoclingDocument.load_from_doctags(doctags_doc, document_name="Document")
                
                page_text = doc_ling.export_to_markdown()
                all_text += page_text + "\n\n"
                
                # Clean up
                os.remove(temp_image_path)
            
            doc.close()
            return all_text.strip()
            
        except Exception as e:
            logger.error(f"SmolDocLing PDF extraction failed: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file with enhanced formatting"""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text with better formatting
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " | "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file with encoding detection"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read().strip()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try with error handling
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read().strip()
                
        except Exception as e:
            logger.error(f"Error reading TXT file: {e}")
            return ""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from various document formats using DocLing for PDFs"""
        file_path = Path(file_path)
        
        if file_path.suffix.lower() == '.pdf':
            return self.extract_text_from_pdf(str(file_path))
        elif file_path.suffix.lower() == '.docx':
            return self.extract_text_from_docx(str(file_path))
        elif file_path.suffix.lower() == '.txt':
            return self.extract_text_from_txt(str(file_path))
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
    
    def preprocess_text(self, text: str) -> str:
        """Enhanced text preprocessing for resume parsing"""
        # Remove extra whitespace but preserve structure
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Normalize paragraph breaks
        text = re.sub(r' +', ' ', text)  # Normalize spaces
        
        # Remove special characters but keep important ones
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\{\}\@\#\$\%\&\+]', '', text)
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        return text.strip()
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split text into overlapping chunks with better boundary detection"""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings
                sentence_endings = ['.', '!', '?', '\n\n']
                break_point = -1
                
                for ending in sentence_endings:
                    pos = chunk.rfind(ending)
                    if pos > start + chunk_size // 2:
                        break_point = max(break_point, pos + len(ending))
                
                if break_point > start + chunk_size // 2:
                    chunk = text[start:break_point]
                    end = break_point
            
            chunks.append(chunk.strip())
            start = end - overlap
            
            if start >= len(text):
                break
        
        return chunks
    
    def extract_structured_data(self, text: str) -> Dict[str, Any]:
        """Enhanced structured data extraction from resume text"""
        try:
            # Extract key information with improved patterns
            structured_data = {
                'personal_info': self._extract_personal_info(text),
                'education': self._extract_education(text),
                'experience': self._extract_experience(text),
                'skills': self._extract_skills(text),
                'contact': self._extract_contact_info(text),
                'summary': self._extract_summary(text),
                'raw_text': text[:2000]  # Increased for better context
            }
            
            return structured_data
            
        except Exception as e:
            logger.error(f"Error in structured data extraction: {e}")
            return {
                'error': str(e),
                'raw_text': text[:2000]
            }
    
    def _extract_personal_info(self, text: str) -> Dict[str, str]:
        """Enhanced personal information extraction"""
        info = {}
        
        # Name extraction (look for patterns)
        lines = text.split('\n')
        for i, line in enumerate(lines[:5]):  # Check first 5 lines
            line = line.strip()
            if line and len(line) < 100 and not any(keyword in line.lower() for keyword in ['email', 'phone', 'linkedin', 'github']):
                # Simple name pattern (2-3 words, no special chars)
                if re.match(r'^[A-Z][a-z]+(\s+[A-Z][a-z]+){1,2}$', line):
                    info['name'] = line
                    break
        
        # Email extraction with better pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            info['email'] = emails[0]
        
        # Phone extraction with multiple formats
        phone_patterns = [
            r'(\+?1?[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})',
            r'\+?[0-9]{1,3}[-.\s]?[0-9]{3,4}[-.\s]?[0-9]{3,4}[-.\s]?[0-9]{3,4}',
            r'[0-9]{3}[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}'
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                if isinstance(phones[0], tuple):
                    info['phone'] = ''.join(phones[0])
                else:
                    info['phone'] = phones[0]
                break
        
        return info
    
    def _extract_education(self, text: str) -> List[Dict[str, str]]:
        """Enhanced education information extraction"""
        education = []
        
        # Look for education section with better keywords
        education_keywords = ['education', 'academic', 'degree', 'university', 'college', 'school', 'bachelor', 'master', 'phd']
        
        lines = text.split('\n')
        in_education_section = False
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            
            # Check if we're entering education section
            if any(keyword in line_lower for keyword in education_keywords):
                in_education_section = True
                education.append({
                    'section': line.strip(),
                    'keyword': 'education'
                })
                continue
            
            # If in education section, collect relevant lines
            if in_education_section and line.strip():
                # Stop if we hit another major section
                if any(section in line_lower for section in ['experience', 'work', 'skills', 'projects']):
                    break
                
                education.append({
                    'section': line.strip(),
                    'keyword': 'education_detail'
                })
        
        return education
    
    def _extract_experience(self, text: str) -> List[Dict[str, str]]:
        """Enhanced work experience extraction"""
        experience = []
        
        # Look for experience section with better keywords
        experience_keywords = ['experience', 'work history', 'employment', 'career', 'professional experience']
        
        lines = text.split('\n')
        in_experience_section = False
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            
            # Check if we're entering experience section
            if any(keyword in line_lower for keyword in experience_keywords):
                in_experience_section = True
                experience.append({
                    'section': line.strip(),
                    'keyword': 'experience'
                })
                continue
            
            # If in experience section, collect relevant lines
            if in_experience_section and line.strip():
                # Stop if we hit another major section
                if any(section in line_lower for section in ['education', 'skills', 'projects', 'certifications']):
                    break
                
                experience.append({
                    'section': line.strip(),
                    'keyword': 'experience_detail'
                })
        
        return experience
    
    def _extract_skills(self, text: str) -> List[str]:
        """Enhanced skills extraction with more comprehensive keyword list"""
        skills = []
        
        # Comprehensive skill keywords
        skill_keywords = [
            # Programming Languages
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'go', 'rust', 'php', 'ruby', 'swift', 'kotlin',
            'scala', 'r', 'matlab', 'perl', 'bash', 'powershell', 'sql',
            
            # Frameworks & Libraries
            'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask', 'spring', 'laravel', 'asp.net',
            'jquery', 'bootstrap', 'tailwind', 'material-ui', 'redux', 'vuex', 'graphql', 'rest',
            
            # Cloud & DevOps
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git', 'github', 'gitlab', 'bitbucket',
            'terraform', 'ansible', 'chef', 'puppet', 'ci/cd', 'devops',
            
            # Databases
            'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'cassandra', 'dynamodb', 'sqlite',
            'oracle', 'sql server', 'firebase', 'supabase',
            
            # Data Science & ML
            'machine learning', 'ai', 'tensorflow', 'pytorch', 'scikit-learn', 'pandas', 'numpy', 'matplotlib',
            'seaborn', 'jupyter', 'spark', 'hadoop', 'kafka', 'airflow',
            
            # Tools & Platforms
            'vscode', 'intellij', 'eclipse', 'vim', 'emacs', 'jira', 'confluence', 'slack', 'teams',
            'zoom', 'figma', 'sketch', 'adobe', 'photoshop', 'illustrator',
            
            # Methodologies
            'agile', 'scrum', 'kanban', 'waterfall', 'tdd', 'bdd', 'lean', 'six sigma'
        ]
        
        text_lower = text.lower()
        for skill in skill_keywords:
            if skill in text_lower:
                skills.append(skill)
        
        return list(set(skills))  # Remove duplicates
    
    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        """Enhanced contact information extraction"""
        contact = {}
        
        # Email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            contact['email'] = emails[0]
        
        # Phone with multiple formats
        phone_patterns = [
            r'(\+?1?[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})',
            r'\+?[0-9]{1,3}[-.\s]?[0-9]{3,4}[-.\s]?[0-9]{3,4}[-.\s]?[0-9]{3,4}',
            r'[0-9]{3}[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}'
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                if isinstance(phones[0], tuple):
                    contact['phone'] = ''.join(phones[0])
                else:
                    contact['phone'] = phones[0]
                break
        
        # LinkedIn
        linkedin_patterns = [
            r'linkedin\.com/in/[\w-]+',
            r'linkedin\.com/company/[\w-]+',
            r'@[\w-]+'  # Twitter-style handles
        ]
        
        for pattern in linkedin_patterns:
            linkedin = re.findall(pattern, text)
            if linkedin:
                contact['linkedin'] = linkedin[0]
                break
        
        # GitHub
        github_pattern = r'github\.com/[\w-]+'
        github = re.findall(github_pattern, text)
        if github:
            contact['github'] = github[0]
        
        return contact
    
    def _extract_summary(self, text: str) -> str:
        """Enhanced summary/objective extraction"""
        summary_keywords = ['summary', 'objective', 'profile', 'about', 'overview']
        
        lines = text.split('\n')
        summary_lines = []
        in_summary_section = False
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Check if we're entering summary section
            if any(keyword in line_lower for keyword in summary_keywords):
                in_summary_section = True
                continue
            
            # If in summary section, collect lines until we hit another section
            if in_summary_section and line.strip():
                if any(section in line_lower for section in ['experience', 'education', 'skills', 'work']):
                    break
                summary_lines.append(line.strip())
        
        return ' '.join(summary_lines) if summary_lines else ""
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Main method to process a document with enhanced error handling"""
        try:
            # Extract text
            raw_text = self.extract_text(file_path)
            if not raw_text:
                return {'error': 'No text extracted from document'}
            
            # Preprocess text
            processed_text = self.preprocess_text(raw_text)
            
            # Extract structured data
            structured_data = self.extract_structured_data(processed_text)
            
            # Add metadata
            structured_data['metadata'] = {
                'file_path': file_path,
                'file_size': os.path.getsize(file_path),
                'text_length': len(processed_text),
                'chunks': len(self.chunk_text(processed_text)),
                'processing_method': 'docling_with_ocr',
                'ocr_enabled': self.use_ocr,
                'ocr_language': self.ocr_language
            }
            
            return structured_data
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            return {'error': str(e)} 